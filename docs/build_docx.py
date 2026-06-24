# -*- coding: utf-8 -*-
"""Convert 论文-偏见调停有效性模拟实验.md to academic DOCX."""
import re, os, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

MD = os.path.join(os.path.dirname(__file__), '论文-偏见调停有效性模拟实验.md')
FIG_DIR = os.path.join(os.path.dirname(__file__), 'figures')
OUT = os.path.join(os.path.dirname(__file__), '论文-偏见调停有效性模拟实验.docx')

# ---- helpers ----
def _font(run, cn='FangSong', en='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rf)
    for k, v in [('eastAsia', cn), ('ascii', en), ('hAnsi', en), ('cs', en)]:
        rf.set(qn(f'w:{k}'), v)

def _spacing(para, before=0, after=0, line=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def _indent(para, pt=24):
    para.paragraph_format.first_line_indent = Pt(pt)

def _clean_text(text):
    text = text.replace('\\"', '"')
    text = re.sub(r'-{3,}', '——', text)
    text = re.sub(r'\$\^\{([^}]+)\}\$', r'\1', text)
    text = re.sub(r'\$\^(\d+)\$', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = re.sub(r'\$\$([^$]+)\$\$', r'\1', text)
    return text

def _render_inline(p, text, cn='FangSong', en='Times New Roman', sz=12, indent=True):
    """Parse inline markdown: **bold**, *italic*, [text](url), ![alt](url)."""
    text = _clean_text(text)
    # strip images (they're rendered separately)
    text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', text)
    # strip link URLs, keep text
    text = re.sub(r'\[([^\]]*)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]*)\]\[[^\]]*\]', r'\1', text)

    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            r = p.add_run(text[last:m.start()])
            _font(r, cn, en, sz)
        if m.group(2):   # **bold**
            r = p.add_run(m.group(2))
            _font(r, cn, en, sz, bold=True)
        elif m.group(3): # *italic*
            r = p.add_run(m.group(3))
            _font(r, cn, en, sz, italic=True)
        elif m.group(4): # `code`
            r = p.add_run(m.group(4))
            _font(r, cn, en, sz)
        last = m.end()
    if last < len(text):
        r = p.add_run(text[last:])
        _font(r, cn, en, sz)

# ---- pre-build: check images are valid ----
from docx.image.image import Image as DocxImage
from PIL import Image as PILImage

fig_files = [f for f in os.listdir(FIG_DIR) if f.endswith(('.png','.jpg','.webp','.jpeg'))]
for f in fig_files:
    fp = os.path.join(FIG_DIR, f)
    try:
        DocxImage.from_file(fp)
    except Exception:
        out = os.path.splitext(fp)[0] + '.png'
        PILImage.open(fp).save(out)
        print(f'  Converted: {f} -> {os.path.basename(out)}')
        os.remove(fp)

# ---- main build ----
doc = Document()

# Wipe Normal style pPr/rPr
ns_n = doc.styles['Normal']
for tag in (qn('w:pPr'), qn('w:rPr')):
    el = ns_n.element.find(tag)
    if el is not None: ns_n.element.remove(el)

# Page setup
for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)
    # Page number
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run()
    _font(fr, 'Times New Roman', 'Times New Roman', size=9)
    fr._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    fr._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    fr._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

# ---- parse markdown ----
with open(MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Parse into typed blocks
blocks = []
i = 0
in_table = False
table_rows = []
table_has_header = False

while i < len(lines):
    line = lines[i].rstrip()

    # Skip horizontal rules
    if re.match(r'^-{3,}$', line.strip()):
        i += 1
        continue

    # Table detection
    if line.strip().startswith('|') and '|' in line[1:]:
        if not in_table:
            in_table = True
            table_rows = []
            table_has_header = False
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        table_rows.append(cells)
        i += 1
        # Peek ahead for separator row
        if i < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i].strip()):
            table_has_header = True
            i += 1  # skip separator
        continue
    elif in_table:
        # End table
        blocks.append(('table', table_rows, table_has_header))
        in_table = False
        table_rows = []
        table_has_header = False
        continue

    # Image: ![图N](path)
    img_match = re.match(r'^!\[(图\d+)\]\((.+)\)$', line.strip())
    if img_match:
        alt = img_match.group(1)
        img_path = img_match.group(2)
        # Check next line is bold caption
        caption = ''
        if i + 1 < len(lines) and re.match(r'^\*\*图\d+', lines[i+1].strip()):
            caption = re.sub(r'\*\*', '', lines[i+1].strip())
            i += 1
        blocks.append(('image', alt, img_path, caption))
        i += 1
        continue

    # Blank line
    if not line.strip():
        i += 1
        continue

    # Collect paragraph text (may span multiple lines)
    para_lines = [line]
    i += 1
    while i < len(lines) and lines[i].strip() and \
          not lines[i].strip().startswith('|') and \
          not re.match(r'^!\[图\d+\]', lines[i].strip()) and \
          not re.match(r'^#{1,4}\s', lines[i].strip()) and \
          not re.match(r'^-{3,}$', lines[i].strip()):
        para_lines.append(lines[i].rstrip())
        i += 1
    full_text = ' '.join(para_lines).strip()

    # Classify
    if re.match(r'^# [^#]', full_text):
        blocks.append(('title', full_text[2:]))
    elif re.match(r'^## (.+)', full_text):
        blocks.append(('h1', re.match(r'^## (.+)', full_text).group(1)))
    elif re.match(r'^### (.+)', full_text):
        blocks.append(('h2', re.match(r'^### (.+)', full_text).group(1)))
    elif re.match(r'^\*\*摘\s*要', full_text):
        # Strip ** markers and extract content after 摘要：
        clean = re.sub(r'\*\*', '', full_text)
        clean = re.sub(r'^摘\s*要[：:]?\s*', '', clean)
        blocks.append(('abstract', clean))
    elif re.match(r'^\*\*关键词', full_text):
        kw = re.sub(r'^\*\*关键词[：:]\s*\*\*', '', full_text).strip()
        blocks.append(('keywords', kw))
    elif re.match(r'^\*\*图\d+\s+', full_text):
        # Standalone figure caption without image above
        blocks.append(('fig_caption', re.sub(r'\*\*', '', full_text)))
    elif re.match(r'^\*\*表\d+', full_text):
        blocks.append(('table_caption', re.sub(r'\*\*', '', full_text)))
    elif re.match(r'^[（\(]实验系统基于', full_text):
        blocks.append(('appendix', full_text.strip('（').strip('）')))
    else:
        blocks.append(('body', full_text))

# Handle trailing table
if in_table and table_rows:
    blocks.append(('table', table_rows, table_has_header))

# ---- render blocks ----
for b in blocks:
    p = doc.add_paragraph()
    _spacing(p, before=0, after=0, line=1.5)

    if b[0] == 'title':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(b[1])
        _font(r, 'STXinwei', 'Times New Roman', size=20)
        # Add blank line after title
        p2 = doc.add_paragraph()
        _spacing(p2, before=0, after=0, line=1.5)

    elif b[0] == 'abstract':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _indent(p, 24)
        r_label = p.add_run('摘　要：')
        _font(r_label, 'FangSong', 'Times New Roman', size=12, bold=True)
        _render_inline(p, b[1], 'FangSong', 'Times New Roman', 12)

    elif b[0] == 'keywords':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _indent(p, 24)
        r = p.add_run('关键词：')
        _font(r, 'FangSong', 'Times New Roman', size=12, bold=True)
        _render_inline(p, b[1], 'FangSong', 'Times New Roman', 12)

    elif b[0] == 'h1':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(b[1])
        _font(r, 'SimHei', 'Times New Roman', size=14, bold=True)

    elif b[0] == 'h2':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(b[1])
        _font(r, 'KaiTi', 'Times New Roman', size=14, bold=True)

    elif b[0] == 'body':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _indent(p, 24)
        _render_inline(p, b[1], 'FangSong', 'Times New Roman', 12)

    elif b[0] == 'table_caption':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p, before=6, after=0, line=1.5)
        r = p.add_run(b[1])
        _font(r, 'FangSong', 'Times New Roman', size=10.5, bold=True)

    elif b[0] == 'fig_caption':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=0, after=4, line=1.5)
        r = p.add_run(b[1])
        _font(r, 'FangSong', 'Times New Roman', size=10.5, bold=True)

    elif b[0] == 'image':
        alt, img_path, caption = b[1], b[2], b[3]
        # Image paragraph
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p_img, before=6, after=0, line=1.5)
        full_path = os.path.join(os.path.dirname(__file__), img_path)
        if os.path.exists(full_path):
            r_img = p_img.add_run()
            r_img.add_picture(full_path, width=Inches(5.5))
        else:
            r_img = p_img.add_run(f'[图片缺失: {img_path}]')
            _font(r_img, 'FangSong', 'Times New Roman', 10)
        # Caption paragraph
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p_cap, before=0, after=4, line=1.5)
            r_cap = p_cap.add_run(caption)
            _font(r_cap, 'FangSong', 'Times New Roman', size=10.5, bold=True)

    elif b[0] == 'table':
        rows, has_header = b[1], b[2]
        ncols = max(len(r) for r in rows)
        nrows = len(rows)
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = 'Table Grid'

        usable = Cm(14.64)  # A4 usable width after margins
        col_w = int(usable / ncols)

        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                cell = table.cell(ri, ci)
                cell.paragraphs[0].clear()
                _spacing(cell.paragraphs[0], before=0, after=0, line=1.15)
                r = cell.paragraphs[0].add_run(cell_text)
                sz = 10 if (has_header and ri == 0) or ncols >= 5 else 12
                _font(r, 'FangSong', 'Times New Roman', size=sz, bold=(has_header and ri == 0))
            # Pad short rows
            for ci in range(len(row_data), ncols):
                cell = table.cell(ri, ci)
                _spacing(cell.paragraphs[0], before=0, after=0, line=1.15)

        # Blank paragraph after table
        p_after = doc.add_paragraph()
        _spacing(p_after, before=0, after=0, line=1.5)

    elif b[0] == 'appendix':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _indent(p, 24)
        r = p.add_run(b[1])
        _font(r, 'FangSong', 'Times New Roman', size=10.5)

doc.save(OUT)

# ---- post-process: inject docDefaults font sizes ----
import zipfile, shutil
tmp = OUT + '.tmp'
with zipfile.ZipFile(OUT, 'r') as zin:
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/styles.xml':
                tree = etree.fromstring(data)
                ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                # Set docDefaults
                dd = tree.find(f'{ns_w}docDefaults')
                if dd is None:
                    dd = parse_xml(f'<w:docDefaults {nsdecls("w")} />')
                    tree.insert(0, dd)
                # rPrDefault
                rpd = dd.find(f'{ns_w}rPrDefault')
                if rpd is None:
                    rpd = parse_xml(f'<w:rPrDefault {nsdecls("w")} />')
                    dd.append(rpd)
                rPr_dd = rpd.find(f'{ns_w}rPr')
                if rPr_dd is None:
                    rPr_dd = parse_xml(f'<w:rPr {nsdecls("w")} />')
                    rpd.append(rPr_dd)
                for tag, val in [('sz', '21'), ('szCs', '21')]:
                    el = rPr_dd.find(f'{ns_w}{tag}')
                    if el is not None: rPr_dd.remove(el)
                    ne = parse_xml(f'<w:{tag} {nsdecls("w")} w:val="{val}"/>')
                    rPr_dd.append(ne)
                data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
            zout.writestr(item, data)
shutil.move(tmp, OUT)

print(f'Done: {OUT}')
